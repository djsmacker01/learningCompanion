"""
Document Processing Utility
Extract text content from various document formats including PDF, Word, Text, and more
"""

import os
import io
import uuid
from typing import Dict, List, Optional, Tuple
from werkzeug.utils import secure_filename

# PDF Processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    import fitz  # PyMuPDF
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Word Document Processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import win32com.client
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

# Text and Other Formats
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

class DocumentProcessor:
    """Multi-format document content extraction and processing utility"""
    
    def __init__(self):
        self.upload_dir = "uploads"
        self._ensure_upload_dir()
    
    def _ensure_upload_dir(self):
        """Ensure upload directory exists"""
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir)
    
    def process_document(self, file) -> Dict:
        """Complete document processing pipeline for any supported format"""
        try:
            # Save the file
            save_result = self.save_uploaded_file(file)
            if not save_result['success']:
                return save_result
            
            file_type = save_result['file_type']
            file_path = save_result['file_path']
            
            # Route to appropriate processor based on file type
            if file_type == 'pdf':
                return self._process_pdf(file_path, save_result)
            elif file_type in ['docx']:
                return self._process_docx(file_path, save_result)
            elif file_type in ['doc']:
                return self._process_doc(file_path, save_result)
            elif file_type in ['txt', 'text']:
                return self._process_text(file_path, save_result)
            elif file_type in ['md', 'markdown']:
                return self._process_markdown(file_path, save_result)
            elif file_type in ['rtf']:
                return self._process_rtf(file_path, save_result)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {file_type}',
                    'file_info': save_result,
                    'supported_formats': ['pdf', 'docx', 'doc', 'txt', 'md', 'rtf']
                }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }
    
    def _process_pdf(self, file_path: str, file_info: Dict) -> Dict:
        """Process PDF files using the existing PDF processor"""
        try:
            from app.utils.pdf_processor import PDFProcessor
            pdf_processor = PDFProcessor()
            
            # Use existing PDF processing logic
            extraction_result = pdf_processor.extract_text_from_pdf(file_path)
            
            return {
                'success': extraction_result['success'],
                'file_info': file_info,
                'content_info': extraction_result,
                'file_type': 'pdf'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'PDF processing failed: {str(e)}',
                'file_info': file_info,
                'file_type': 'pdf',
                'extraction_method': 'error'
            }
    
    def _process_docx(self, file_path: str, file_info: Dict) -> Dict:
        """Process DOCX files with better formatting preservation"""
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'DOCX processing not available. Please install python-docx library.',
                'file_info': file_info,
                'suggestions': ['Install python-docx: pip install python-docx']
            }
        
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs with better formatting
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    # Check if this looks like a heading (shorter, might be bold)
                    if len(paragraph_text) < 100 and not paragraph_text.endswith('.'):
                        # Likely a heading, add extra spacing
                        text_content += "\n" + paragraph_text + "\n"
                    else:
                        text_content += paragraph_text + "\n"
            
            # Add separator before tables
            if doc.tables:
                text_content += "\n" + "="*50 + "\n"
            
            # Extract text from tables with better formatting
            for table_idx, table in enumerate(doc.tables):
                text_content += f"\nTABLE {table_idx + 1}:\n"
                for row_idx, row in enumerate(table.rows):
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
                text_content += "\n"
            
            # Clean and process text
            processed_content = self._process_extracted_text(text_content, "docx")
            
            return {
                'success': True,
                'file_info': file_info,
                'content_info': processed_content,
                'file_type': 'docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing DOCX file: {str(e)}',
                'file_info': file_info
            }
    
    def _process_doc(self, file_path: str, file_info: Dict) -> Dict:
        """Process legacy DOC files using COM (Windows only)"""
        if not DOC_AVAILABLE:
            return {
                'success': False,
                'error': 'DOC processing not available. Legacy DOC files require Windows COM support.',
                'file_info': file_info,
                'suggestions': [
                    'Convert DOC to DOCX format before uploading',
                    'Use Windows system with python-win32 installed'
                ]
            }
        
        try:
            # Use COM to open Word and extract text
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(file_path)
            text_content = doc.Content.Text
            doc.Close()
            word.Quit()
            
            # Clean and process text
            processed_content = self._process_extracted_text(text_content, "doc")
            
            return {
                'success': True,
                'file_info': file_info,
                'content_info': processed_content,
                'file_type': 'doc'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing DOC file: {str(e)}',
                'file_info': file_info,
                'suggestions': [
                    'Try converting to DOCX format',
                    'Ensure Word is installed on the system'
                ]
            }
    
    def _process_text(self, file_path: str, file_info: Dict) -> Dict:
        """Process plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252', 'ascii']
            text_content = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                return {
                    'success': False,
                    'error': 'Could not decode text file. Unsupported encoding.',
                    'file_info': file_info
                }
            
            # Clean and process text
            processed_content = self._process_extracted_text(text_content, "text")
            
            return {
                'success': True,
                'file_info': file_info,
                'content_info': processed_content,
                'file_type': 'text'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing text file: {str(e)}',
                'file_info': file_info
            }
    
    def _process_markdown(self, file_path: str, file_info: Dict) -> Dict:
        """Process Markdown files"""
        try:
            # Read the markdown file
            with open(file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # Convert to HTML for better display (optional)
            html_content = ""
            if MARKDOWN_AVAILABLE:
                html_content = markdown.markdown(markdown_content)
            
            # Clean and process text (use raw markdown for content)
            processed_content = self._process_extracted_text(markdown_content, "markdown")
            
            # Add HTML content if available
            if html_content:
                processed_content['html_content'] = html_content
            
            return {
                'success': True,
                'file_info': file_info,
                'content_info': processed_content,
                'file_type': 'markdown'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing markdown file: {str(e)}',
                'file_info': file_info
            }
    
    def _process_rtf(self, file_path: str, file_info: Dict) -> Dict:
        """Process RTF files (basic text extraction)"""
        try:
            # Simple RTF text extraction (removes RTF formatting)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            # Basic RTF cleaning (remove formatting codes)
            import re
            # Remove RTF control words and groups
            text_content = re.sub(r'\\[a-z]+\d*\s?', '', rtf_content)
            text_content = re.sub(r'[{}]', '', text_content)
            text_content = re.sub(r'\s+', ' ', text_content)
            
            # Clean and process text
            processed_content = self._process_extracted_text(text_content, "rtf")
            
            return {
                'success': True,
                'file_info': file_info,
                'content_info': processed_content,
                'file_type': 'rtf'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing RTF file: {str(e)}',
                'file_info': file_info
            }
    
    def _process_extracted_text(self, text_content: str, file_type: str) -> Dict:
        """Process and clean extracted text"""
        if not text_content:
            return {
                'cleaned_text': '',
                'word_count': 0,
                'key_sections': [],
                'extraction_method': file_type,
                'file_type': file_type
            }
        
        # Clean the text
        cleaned_text = self._clean_text(text_content)
        
        # Extract key sections
        key_sections = self._extract_key_sections(cleaned_text)
        
        return {
            'cleaned_text': cleaned_text,
            'word_count': len(cleaned_text.split()),
            'key_sections': key_sections,
            'extraction_method': file_type,
            'file_type': file_type
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text while preserving structure"""
        import re
        
        if not text:
            return ""
        
        # Split into lines for better processing
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                # Preserve empty lines for structure
                cleaned_lines.append("")
                continue
            
            # Clean the line but preserve structure
            # Remove excessive spaces but keep single spaces
            line = re.sub(r'\s+', ' ', line)
            
            # Remove special characters that aren't useful but keep punctuation
            line = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\'\"\/\@]+', ' ', line)
            
            # Clean up multiple spaces again
            line = re.sub(r'\s+', ' ', line)
            
            cleaned_lines.append(line.strip())
        
        # Join lines and clean up excessive empty lines
        result = '\n'.join(cleaned_lines)
        
        # Remove more than 2 consecutive empty lines
        result = re.sub(r'\n{3,}', '\n\n', result)
        
        return result.strip()
    
    def _extract_key_sections(self, text: str) -> List[Dict]:
        """Extract key sections from the text"""
        import re
        
        sections = []
        
        # Look for common section patterns
        section_patterns = [
            r'(Chapter\s+\d+[:\s]+[^\n]+)',
            r'(Section\s+\d+[:\s]+[^\n]+)',
            r'(^\d+\.\s+[^\n]+)',  # Numbered sections
            r'(^[A-Z][A-Z\s]{2,}:)',  # All caps headings
            r'(^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:)',  # Title case headings
        ]
        
        for pattern in section_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            for match in matches:
                if len(match.strip()) > 3:
                    sections.append({
                        'title': match.strip(),
                        'type': 'heading'
                    })
        
        # Look for key terms
        words = re.findall(r'\b[A-Z][a-z]+\b', text)
        word_freq = {}
        for word in words:
            if len(word) > 4:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Add frequent terms as key concepts
        for word, freq in sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]:
            if freq > 2:
                sections.append({
                    'title': word,
                    'type': 'concept',
                    'frequency': freq
                })
        
        return sections
    
    def save_uploaded_file(self, file) -> Dict:
        """Save uploaded file and return file info"""
        try:
            # Generate secure filename
            filename = secure_filename(file.filename)
            unique_filename = f"{uuid.uuid4()}_{filename}"
            file_path = os.path.join(self.upload_dir, unique_filename)
            
            # Save file
            file.save(file_path)
            
            # Get file info
            file_size = os.path.getsize(file_path)
            file_extension = filename.split('.')[-1].lower() if '.' in filename else 'unknown'
            
            return {
                'success': True,
                'file_path': file_path,
                'filename': unique_filename,
                'original_filename': file.filename,
                'file_size': file_size,
                'file_type': file_extension
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error saving file: {str(e)}'
            }
    
    @staticmethod
    def is_document_processing_available() -> Dict[str, bool]:
        """Check which document processing libraries are available"""
        return {
            'pdf': PDFPLUMBER_AVAILABLE or PDF_AVAILABLE or PDFMINER_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'doc': DOC_AVAILABLE,
            'markdown': MARKDOWN_AVAILABLE,
            'ocr': OCR_AVAILABLE,
            'text': True,  # Always available
            'rtf': True    # Basic processing always available
        }
    
    @staticmethod
    def get_supported_formats() -> List[str]:
        """Get list of supported file formats"""
        return ['pdf', 'docx', 'doc', 'txt', 'md', 'markdown', 'rtf']
    
    @staticmethod
    def get_required_libraries() -> List[str]:
        """Get list of required libraries for full functionality"""
        libraries = []
        if not PDFPLUMBER_AVAILABLE:
            libraries.append('pdfplumber')
        if not DOCX_AVAILABLE:
            libraries.append('python-docx')
        if not MARKDOWN_AVAILABLE:
            libraries.append('markdown')
        if not OCR_AVAILABLE:
            libraries.append('pytesseract')
            libraries.append('PyMuPDF')
        return libraries
